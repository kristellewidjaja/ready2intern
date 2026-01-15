"""
Resume parsing service for extracting text from PDF and DOCX files.
"""

import logging
from pathlib import Path
from typing import Tuple
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """Service for parsing resume files and extracting text content."""
    
    def __init__(self):
        """Initialize resume parser."""
        logger.info("ResumeParser initialized")
    
    def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text content from resume file.
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Tuple of (extracted_text, error_message)
            If successful, error_message is empty string
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        try:
            if file_ext == ".pdf":
                return self._extract_from_pdf(path), ""
            elif file_ext == ".docx":
                return self._extract_from_docx(path), ""
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return "", str(e)
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        logger.info(f"Extracting text from PDF: {file_path}")
        
        text_content = []
        
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF has {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
            result = "\n".join(text_content)
            logger.info(f"Extracted {len(result)} characters from PDF")
            
            if not result.strip():
                raise ValueError("PDF appears to be empty or contains only images")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}") from e
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            result = "\n".join(text_content)
            logger.info(f"Extracted {len(result)} characters from DOCX")
            
            if not result.strip():
                raise ValueError("DOCX appears to be empty")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}") from e
